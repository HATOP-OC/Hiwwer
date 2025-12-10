from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_styles(doc):
    """Налаштування базових стилів документа"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.first_line_indent = Cm(1.25)
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_cell_shading(cell, color):
    """Додає заливку комірки таблиці"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_number(run):
    """Додає номер сторінки"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_footer_numbering(doc, start_from=2):
    """Налаштування нумерації сторінок у футері"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(0)
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    add_page_number(run)

def set_first_page_different(doc):
    """Встановлює різний футер для першої сторінки (без номера)"""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

def add_heading(doc, text, level=1, add_page_break=True):
    """Додає заголовок розділу або підрозділу"""
    if level == 1 and add_page_break:
        if len(doc.paragraphs) > 1:
            doc.add_page_break()
        
    heading = doc.add_heading(text.upper() if level == 1 else text, level=level)
    
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.first_line_indent = Cm(0)
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.first_line_indent = Cm(1.25)
    
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        rFonts = run._element.rPr.rFonts if run._element.rPr is not None else run._element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    heading.paragraph_format.keep_with_next = True

def add_subheading(doc, text, level=3):
    """Додає підзаголовок третього рівня (3.3. 1, 3.3.2 і т.д.)"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    heading.paragraph_format.first_line_indent = Cm(1.25)
    heading.paragraph_format.keep_with_next = True

def add_text(doc, text):
    """Додає звичайний абзац тексту"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def add_list_item(doc, text):
    """Додає елемент ненумерованого списку"""
    p = doc.add_paragraph("– " + text)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def add_code_snippet(doc, code):
    """Додає блок коду"""
    lines = code.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

def add_table_of_contents(doc):
    """Додає сторінку ЗМІСТ"""
    doc.add_page_break()
    
    # Заголовок ЗМІСТ
    p = doc.add_paragraph("ЗМІСТ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Елементи змісту
    toc_items = [
        ("ВСТУП", "3"),
        ("РОЗДІЛ 1.  АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ ТА ВИМОГ ДО БАЗИ ДАНИХ МАРКЕТПЛЕЙСУ", "5"),
        ("1.1.  Дослідження типів маркетплейсів та їх функціональних особливостей", "5"),
        ("1.2.  Вимоги до структури даних, цілісності та безпеки", "7"),
        ("1.3. Особливості інтеграції з Telegram та системами штучного інтелекту", "9"),
        ("РОЗДІЛ 2.  ПРОЕКТУВАННЯ АРХІТЕКТУРИ БАЗИ ДАНИХ ТА ВИБІР ІНСТРУМЕНТАРІЮ", "11"),
        ("2.1. Обґрунтування вибору СУБД PostgreSQL", "11"),
        ("2.2. Концептуальне проектування та ER-діаграма", "13"),
        ("2.3. Логічне проектування та нормалізація", "15"),
        ("РОЗДІЛ 3. РЕАЛІЗАЦІЯ БАЗИ ДАНИХ СИСТЕМИ «HIWWER»", "17"),
        ("3.1.  Фізична реалізація структури даних", "17"),
        ("3.2. Розробка серверної логіки (Тригери та Функції)", "19"),
        ("3.3.  Реалізація специфічних модулів та інтеграцій", "21"),
        ("3.3.1. Модуль історії статусів та чатів", "21"),
        ("3.3.2. Багатомовні політики", "22"),
        ("3.3.3.  Інтеграція з Telegram-ботом", "22"),
        ("3.3.4.  Адміністративна панель та управління системою", "23"),
        ("3.3.5. Архітектура взаємодії Telegram-бота з базою даних", "25"),
        ("3.3.6.  Механізм резервного копіювання та міграції", "27"),
        ("ВИСНОВКИ", "28"),
        ("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", "30"),
        ("ДОДАТКИ", "31"),
    ]
    
    # Налаштування табуляції для вирівнювання номерів сторінок
    tab_stops = doc.styles['Normal'].paragraph_format.tab_stops
    # Видаляємо старі табулятори, якщо є
    for tab in tab_stops:
        del tab
    # Додаємо табулятор праворуч (приблизно 16 см для А4 з полями)
    # 21cm (width) - 3cm (left) - 1cm (right) = 17cm
    tab_stops.add_tab_stop(Cm(16.5), alignment=WD_TABLE_ALIGNMENT.RIGHT, leader=WD_TABLE_ALIGNMENT.LEFT) # leader=1 is dots

    for item_text, page_num in toc_items:
        p = doc.add_paragraph()
        
        # Визначаємо рівень вкладеності та відступи
        if item_text.startswith("РОЗДІЛ") or item_text in ["ВСТУП", "ВИСНОВКИ", "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", "ДОДАТКИ"]:
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(item_text)
            run.font.bold = True
        elif item_text[0].isdigit() and "." in item_text and item_text.count(".") == 1:
            p.paragraph_format.first_line_indent = Cm(0.5)
            run = p.add_run(item_text)
        elif item_text[0].isdigit() and item_text.count(". ") >= 2:
            p.paragraph_format.first_line_indent = Cm(1.0)
            run = p.add_run(item_text)
        else:
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(item_text)
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Додаємо табуляцію та номер сторінки
        # Використовуємо \t для табуляції
        run_tab = p.add_run("\t" + page_num)
        run_tab.font.name = 'Times New Roman'
        run_tab.font.size = Pt(14)
        
        # Налаштування абзацу
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_title_page(doc):
    """Створює титульну сторінку"""
    # Назва міністерства
    p = doc.add_paragraph("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Назва університету (якщо потрібно)
    p = doc.add_paragraph("НАЦІОНАЛЬНИЙ ТЕХНІЧНИЙ УНІВЕРСИТЕТ УКРАЇНИ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph("«КИЇВСЬКИЙ ПОЛІТЕХНІЧНИЙ ІНСТИТУТ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph("ІМЕНІ ІГОРЯ СІКОРСЬКОГО»")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Факультет та кафедра
    p = doc.add_paragraph("Факультет інформатики та обчислювальної техніки")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph("Кафедра обчислювальної техніки")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(48)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # КУРСОВА РОБОТА
    p = doc.add_paragraph("КУРСОВА РОБОТА")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.font.bold = True
    
    # Дисципліна
    p = doc.add_paragraph("з дисципліни «Організація баз даних»")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Тема
    p = doc.add_paragraph("на тему:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph("«Проектування та реалізація реляційної бази даних для маркетплейсу цифрових послуг на основі PostgreSQL з інтеграцією в Telegram бота та з ШІ асистентом»")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(72)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Інформація про виконавця (права сторона)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Виконав:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("студент ІІІ курсу, групи ІО-з31")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("________________ Прізвище І.Б.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Керівник:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("к.т.н., доцент")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(72)
    run = p.add_run("________________ Прізвище І.Б.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Місто та рік
    p = doc.add_paragraph("Київ – 2025")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

def create_coursework():
    """Головна функція створення курсової роботи"""
    try:
        doc = Document('титулка.docx')
    except:
        doc = Document()
    
    # Налаштування секцій та полів
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.0)

    setup_styles(doc)
    set_first_page_different(doc)
    add_footer_numbering(doc)

    # --- ТИТУЛЬНА СТОРІНКА ---
    # add_title_page(doc)

    # --- ЗМІСТ ---
    add_table_of_contents(doc)

    # --- ВСТУП ---
    add_heading(doc, "ВСТУП", 1)
    
    add_text(doc, "Актуальність теми дослідження.  В умовах глобальної цифровізації економіки та стрімкого розвитку інформаційних технологій спостерігається фундаментальна трансформація ринку праці. Традиційні моделі зайнятості, що базувалися на довгострокових трудових договорах та фіксованому робочому графіку, поступово поступаються місцем гнучким формам співпраці. Серед них провідне місце займає фріланс та гіг-економіка.  Електронні маркетплейси цифрових послуг стають ключовими інфраструктурними елементами цієї нової економіки, забезпечуючи ефективну, швидку та прозору взаємодію між замовниками та виконавцями незалежно від їх географічного розташування та часових поясів.")
    add_text(doc, "Згідно з останніми дослідженнями, обсяг ринку фріланс-послуг зростає щорічно на 15-20%, що створює попит на надійні технологічні платформи. Сучасний маркетплейс — це вже не просто дошка оголошень, а складна високонавантажена інформаційна система, яка повинна гарантувати безпеку фінансових транзакцій, захист інтелектуальної власності, прозорість рейтингів та зручність комунікації. Зростання обсягів даних, які генеруються такими системами (профілі, портфоліо, історія замовлень, логи транзакцій, повідомлення), висуває підвищені вимоги до архітектури баз даних.  Від ефективності проектування схеми даних залежить швидкодія системи, її здатність до масштабування, стійкість до пікових навантажень та можливість подальшого розвитку.")
    add_text(doc, "Особливої актуальності в сучасних умовах набуває інтеграція веб-платформ із популярними месенджерами та системами штучного інтелекту. Користувачі очікують миттєвого доступу до інформації через звичні канали комунікації, такі як Telegram, що вимагає реалізації складних механізмів синхронізації даних між веб-сервером та бот-платформою. Впровадження AI-асистентів (на базі LLM, таких як Google Gemini або OpenAI GPT) для автоматизації підтримки, генерації описів завдань та підбору послуг створює нові виклики щодо зберігання та обробки неструктурованих текстових даних, векторних представлень та контексту діалогів.")
    add_text(doc, "Таким чином, розробка ефективної бази даних для маркетплейсу є актуальним науково-практичним завданням, яке вимагає глибокого розуміння теорії реляційних баз даних, сучасних методів обробки інформації та специфіки предметної області.")
    add_text(doc, "Метою курсової роботи є проектування та реалізація реляційної бази даних для маркетплейсу цифрових послуг «Hiwwer», яка забезпечує надійне зберігання даних, підтримку складної бізнес-логіки, високу продуктивність, а також безшовну інтеграцію з зовнішніми сервісами Telegram та Google Gemini.")
    add_text(doc, "Для досягнення поставленої мети необхідно вирішити такі завдання:")
    add_list_item(doc, "Провести глибокий аналіз предметної області, дослідити існуючі аналоги (Upwork, Fiverr, Freelancehunt) та сформулювати детальні вимоги до функціоналу та структури даних.")
    add_list_item(doc, "Обґрунтувати вибір технологічного стеку, зокрема СУБД PostgreSQL, провівши порівняльний аналіз з альтернативними рішеннями (MySQL, MongoDB, Oracle).")
    add_list_item(doc, "Розробити концептуальну модель бази даних, визначити основні сутності, їх атрибути та типи зв'язків між ними, побудувати ER-діаграму.")
    add_list_item(doc, "Спроектувати логічну схему бази даних та провести її нормалізацію до третьої нормальної форми (3НФ) для забезпечення цілісності даних та усунення надлишковості.")
    add_list_item(doc, "Реалізувати фізичну модель бази даних, використовуючи специфічні можливості PostgreSQL (UUID, JSONB, індекси, партиціювання).")
    add_list_item(doc, "Розробити серверну логіку на рівні бази даних (тригери, збережені процедури та функції) для автоматизації бізнес-процесів та забезпечення цілісності даних.")
    add_list_item(doc, "Реалізувати механізми інтеграції з Telegram-ботом (зв'язування акаунтів, черги повідомлень) та модулем штучного інтелекту (зберігання історії чатів).")
    add_list_item(doc, "Забезпечити надійність та безпеку даних через налаштування прав доступу, резервного копіювання та шифрування чутливої інформації.")
    
    add_text(doc, "Об'єктом дослідження є процеси обробки, зберігання та управління даними в розподілених інформаційних системах електронної комерції.")
    add_text(doc, "Предметом дослідження є методи та засоби проектування реляційної бази даних для маркетплейсу послуг з підтримкою омніканальної взаємодії та елементів штучного інтелекту.")
    add_text(doc, "Методи дослідження.  У роботі використано методи системного аналізу для визначення вимог до системи, методи теорії реляційних баз даних для проектування схеми, методи об'єктно-орієнтованого програмування для реалізації взаємодії з БД, а також методи тестування програмного забезпечення.")

    # --- РОЗДІЛ 1 ---
    add_heading(doc, "РОЗДІЛ 1.  АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ ТА ВИМОГ ДО БАЗИ ДАНИХ МАРКЕТПЛЕЙСУ", 1)
    
    add_heading(doc, "1.1. Дослідження типів маркетплейсів та їх функціональних особливостей", 2)
    add_text(doc, "Ринок електронної комерції послуг характеризується високою різноманітністю бізнес-моделей, кожна з яких висуває свої унікальні вимоги до архітектури даних. Можна виділити кілька основних типів платформ за характером взаємодії учасників:")
    add_list_item(doc, "C2C (Consumer-to-Consumer): платформи, де фізичні особи надають послуги іншим фізичним особам (наприклад, OLX, Kabanchik).  Тут важлива простота реєстрації та швидкість модерації.")
    add_list_item(doc, "B2C (Business-to-Consumer): компанії пропонують послуги кінцевим споживачам.  Вимоги до даних включають підтримку юридичних осіб, податкової звітності та складних каталогів.")
    add_list_item(doc, "B2B (Business-to-Business): платформи для взаємодії бізнесів.  Тут критичними є документообіг, довгострокові контракти та специфічні платіжні методи.")
    add_text(doc, "За моделлю монетизації маркетплейси поділяються на:")
    add_list_item(doc, "Комісійна модель (Commission): платформа стягує відсоток з кожної транзакції (Upwork, Uber).  Це вимагає точного обліку фінансових операцій та підтримки транзакційності.")
    add_list_item(doc, "Передплата (Subscription): користувачі платять за доступ до платформи (LinkedIn Premium).  Необхідно зберігати дані про терміни дії підписок та автоматичне подовження.")
    add_list_item(doc, "Плата за лістинг (Listing fee): оплата за розміщення оголошення (Etsy).")
    add_list_item(doc, "Freemium: базовий функціонал безкоштовний, розширений — платний.")
    add_text(doc, "Для проекту «Hiwwer» обрано гібридну модель, яка поєднує елементи магазину послуг (фіксовані пропозиції, як на Fiverr) та біржі (індивідуальні замовлення, як на Upwork).  Це накладає специфічні вимоги до структури бази даних.  На відміну від товарних маркетплейсів, де товар є статичною одиницею, послуга є динамічною сутністю, яка може змінюватися в процесі виконання (зміна термінів, ціни, обсягу робіт).")
    add_text(doc, "Аналіз бізнес-процесів дозволяє виділити ключові інформаційні об'єкти та потоки даних:")
    add_list_item(doc, "Профілі користувачів: система повинна підтримувати єдиний профіль для клієнта та виконавця, зберігаючи при цьому специфічні атрибути для кожної ролі (портфоліо, навички, сертифікати для виконавця; історія замовлень, платіжні дані для клієнта).")
    add_list_item(doc, "Каталог послуг: ієрархічна структура категорій (деревоподібна структура), система тегів та динамічних фільтрів для швидкого пошуку.  Послуга повинна мати гнучкі атрибути (ціна, терміни, опції, пакети послуг).")
    add_list_item(doc, "Угоди (Замовлення): складний об'єкт, що зв'язує клієнта, виконавця та послугу. Життєвий цикл угоди включає стани: ініціація, узгодження умов, резервування коштів (Escrow), виконання, перевірка, завершення, диспут.  Кожен перехід стану повинен логуватися.")
    add_list_item(doc, "Комунікації: внутрішній месенджер з підтримкою вкладених файлів, який є основним інструментом узгодження деталей замовлення.  Важливо зберігати контекст переписки прив'язаним до конкретного замовлення.")
    add_text(doc, "Важливим аспектом є система арбітражу.  У разі виникнення конфлікту адміністратор повинен мати доступ до повної історії взаємодії сторін: листування, зміни статусів, завантажені файли.  Це вимагає реалізації детального логування всіх дій у базі даних (Audit Trail).")
    add_text(doc, "Окрім того, сучасні маркетплейси все частіше інтегрують елементи соціальних мереж, такі як стрічки новин, лайки, коментарі та можливість підписки на улюблених виконавців. Це створює додаткове навантаження на базу даних і вимагає оптимізації запитів для отримання стрічки активності в реальному часі.  В рамках проекту «Hiwwer» ми зосередимося на базовому функціоналі взаємодії, але закладемо архітектурні можливості для майбутнього розширення соціальних функцій.")

    add_heading(doc, "1.2. Вимоги до структури даних, цілісності та безпеки", 2)
    add_text(doc, "Проектування бази даних для фінансової платформи вимагає суворого дотримання стандартів надійності та безпеки.  Основні вимоги можна класифікувати за кількома напрямками.")
    add_text(doc, "Вимоги до цілісності даних (Data Integrity):")
    add_list_item(doc, "Сутнісна цілісність (Entity Integrity): кожен запис у таблиці повинен мати унікальний первинний ключ.  Використання UUID (Universally Unique Identifier) дозволяє уникнути колізій при розподіленому генерації ключів та приховує реальну кількість записів від конкурентів (на відміну від послідовних ID).")
    add_list_item(doc, "Посилальна цілісність (Referential Integrity): зовнішні ключі повинні посилатися лише на існуючі записи. Необхідно визначити правила каскадного оновлення та видалення (CASCADE, SET NULL, RESTRICT) для запобігання появі «сирітських» записів.  Наприклад, при видаленні користувача його повідомлення можуть залишатися (SET NULL), але його активні сесії повинні видалятися (CASCADE).")
    add_list_item(doc, "Доменна цілісність (Domain Integrity): значення атрибутів повинні відповідати визначеним типам даних та обмеженням (CHECK constraints). Наприклад, ціна не може бути від'ємною, email повинен відповідати формату, а рейтинг повинен бути в межах від 0 до 5.")
    add_text(doc, "Вимоги до транзакційності (ACID):")
    add_text(doc, "Система повинна гарантувати атомарність фінансових операцій. Якщо списання коштів з рахунку клієнта пройшло успішно, а зарахування на транзитний рахунок (Escrow) — ні, вся транзакція повинна бути скасована (ROLLBACK).  PostgreSQL забезпечує це на рівні механізму транзакцій.")
    add_text(doc, "Атомарність (Atomicity) гарантує, що транзакція буде виконана повністю або не виконана зовсім. Узгодженість (Consistency) забезпечує перехід бази даних з одного узгодженого стану в інший, не порушуючи обмежень цілісності.  Ізольованість (Isolation) гарантує, що паралельні транзакції не впливатимуть одна на одну (використання рівнів ізоляції Read Committed або Serializable).  Довговічність (Durability) означає, що результати зафіксованої транзакції будуть збережені навіть у разі збою системи (завдяки Write-Ahead Logging).")
    add_text(doc, "Вимоги до безпеки та відповідності стандартам:")
    add_text(doc, "Захист персональних даних є пріоритетом, особливо в контексті GDPR (General Data Protection Regulation).  Це вимагає мінімізації збережених даних, права на забуття (видалення даних) та захисту від витоків.  Паролі повинні зберігатися виключно у вигляді хешів з сіллю (наприклад, алгоритм bcrypt або Argon2).  Доступ до бази даних повинен бути розмежований на рівні користувачів СУБД (принцип найменших привілеїв). Критичні дані (API ключі, токени) не повинні передаватися у відкритому вигляді.")
    add_text(doc, "Також важливо враховувати принципи CAP-теореми (Consistency, Availability, Partition tolerance). Для фінансових систем пріоритетом є узгодженість (Consistency) та стійкість до розділення (Partition tolerance), навіть якщо це іноді йде на шкоду доступності (Availability).PostgreSQL, як CP-система (в певних конфігураціях), ідеально підходить для таких завдань, забезпечуючи строгу узгодженість даних.")

    add_heading(doc, "1.3. Особливості інтеграції з Telegram та системами штучного інтелекту", 2)
    add_text(doc, "Інтеграція з зовнішніми системами суттєво впливає на архітектуру бази даних та вимагає специфічних рішень.")
    add_text(doc, "Telegram Bot API працює на основі ідентифікаторів `chat_id` та `user_id`.  Ці ідентифікатори є стабільними для конкретного бота, але можуть змінюватися, якщо користувач взаємодіє з іншим ботом.  Для того щоб зв'язати веб-користувача з його Telegram-акаунтом, необхідно реалізувати механізм «рукостискання» (handshake).  Це вимагає створення тимчасових таблиць для зберігання кодів верифікації (OTP - One Time Password) та додавання відповідних полів до основної таблиці користувачів.  Важливо також зберігати стан діалогу (FSM - Finite State Machine) в базі даних, щоб забезпечити безперервність взаємодії при перезапуску бота.")
    add_text(doc, "Крім того, бот повинен мати можливість надсилати сповіщення асинхронно.  Це означає, що база даних повинна виступати як черга повідомлень або джерело подій для бота. Необхідно зберігати статус відправки сповіщень (pending, sent, failed), щоб уникнути дублювання або втрати інформації.  Використання механізму Webhooks для отримання оновлень від Telegram є більш ефективним, ніж Long Polling, але вимагає надійної обробки вхідних запитів та їх збереження в БД.")
    add_text(doc, "Інтеграція з AI (Google Gemini) вимагає зберігання контексту.  Штучний інтелект не має власної пам'яті про попередні запити користувача в рамках сесії (stateless), тому ця інформація повинна передаватися з кожним новим запитом.  Це вимагає створення структури даних для зберігання історії діалогів (`chat_history`), де кожен запис містить роль учасника ('user', 'model', 'system'), текст повідомлення, метадані (кількість токенів) та часову мітку.  Враховуючи потенційно великий обсяг тексту та обмеження контекстного вікна моделі, необхідно оптимізувати зберігання таких даних, можливо, використовуючи стиснення або архівування старих діалогів, а також реалізувати механізм «ковзного вікна» (sliding window) для вибірки останніх повідомлень.")
    add_text(doc, "Важливим аспектом є також кешування відповідей AI.  Оскільки запити до API платних моделей можуть бути дорогими та повільними (латентність може досягати кількох секунд), доцільно зберігати часто запитувані відповіді в базі даних або швидкому кеші (наприклад, Redis). У рамках даної курсової роботи ми розглянемо реалізацію кешування семантично схожих запитів засобами PostgreSQL (використовуючи розширення `pgvector` для векторного пошуку, якщо це необхідно, або просте текстове співпадіння).")

    # --- РОЗДІЛ 2 ---
    add_heading(doc, "РОЗДІЛ 2.  ПРОЕКТУВАННЯ АРХІТЕКТУРИ БАЗИ ДАНИХ ТА ВИБІР ІНСТРУМЕНТАРІЮ", 1)
    
    add_heading(doc, "2.1.  Обґрунтування вибору СУБД PostgreSQL", 2)
    add_text(doc, "Вибір системи управління базами даних (СУБД) є критичним рішенням, яке впливає на весь життєвий цикл проекту. На ринку існує багато рішень, кожне з яких має свої переваги та недоліки. Для даного проекту було проведено порівняльний аналіз найпопулярніших рішень: MySQL, PostgreSQL, MongoDB та Oracle.")
    add_text(doc, "MySQL є найпопулярнішою open-source СУБД для веб-проектів.  Вона проста в налаштуванні та має високу швидкість читання. Проте, MySQL має певні обмеження при роботі зі складними аналітичними запитами, вкладеними транзакціями та JSON-даними порівняно з PostgreSQL.  Її оптимізатор запитів історично поступається PostgreSQL у складності планів виконання.")
    add_text(doc, "MongoDB (NoSQL) забезпечує високу гнучкість схеми (schema-less) та горизонтальну масштабованість. Це чудовий вибір для зберігання неструктурованих даних або логів. Однак, MongoDB не гарантує суворої цілісності даних та ACID-транзакційності на рівні декількох документів (хоча останні версії покращили це), що є критичним ризиком для фінансової системи, де точність балансів є обов'язковою.")
    add_text(doc, "Oracle Database — це потужне комерційне рішення для корпоративного сектору.  Вона має неперевершену надійність та функціональність, але висока вартість ліцензування та складність адміністрування роблять її недоцільною для стартапу або курсового проекту.")
    add_text(doc, "PostgreSQL було обрано як оптимальне рішення з наступних причин:")
    add_list_item(doc, "Об'єктно-реляційна природа: PostgreSQL підтримує не тільки реляційні дані, але й об'єкти, спадкування та користувацькі типи даних.")
    add_list_item(doc, "Підтримка JSONB: це дозволяє зберігати неструктуровані дані (наприклад, параметри замовлення, які відрізняються для різних категорій послуг, або метадані файлів) в ефективному бінарному форматі з можливістю індексації (GIN індекси).  Це поєднує переваги SQL та NoSQL в одній системі.")
    add_list_item(doc, "Розширюваність: PostgreSQL дозволяє підключати розширення, такі як `uuid-ossp` для генерації UUID, `pgcrypto` для шифрування, `PostGIS` для геоданих, що значно спрощує розробку.")
    add_list_item(doc, "Потужна процедурна мова PL/pgSQL: дозволяє реалізувати складну бізнес-логіку (тригери, функції) безпосередньо в базі даних, зменшуючи навантаження на сервер додатку та кількість мережевих запитів (round-trips).")
    add_list_item(doc, "Надійність та відповідність стандартам: PostgreSQL відома своєю стабільністю, повною підтримкою ACID та суворою відповідністю стандартам SQL:2011, що полегшує підтримку та міграцію.")
    add_text(doc, "Архітектура PostgreSQL базується на процесах.  Головний процес `postmaster` керує виділенням пам'яті та запускає фонові процеси (logger, checkpointer, writer, wal writer, autovacuum).  Для кожного клієнтського з'єднання створюється окремий процес `backend`. Це забезпечує стабільність: збій в одному процесі не впливає на роботу всієї системи.  Система управління пам'яттю включає `shared buffers` (спільний кеш сторінок даних для всіх процесів) та `work_mem` (пам'ять для сортування та хешування в межах одного запиту).  Правильне налаштування цих параметрів є критичним для продуктивності.")
    add_text(doc, "Механізм MVCC (Multiversion Concurrency Control) дозволяє читати дані, не блокуючи їх запис, і навпаки.  Це досягається шляхом створення нових версій рядків (tuples) при оновленні, замість перезапису старих.  Старі версії залишаються доступними для транзакцій, що почалися раніше.  Згодом вони видаляються процесом `VACUUM`, щоб звільнити місце.")

    add_heading(doc, "2.2. Концептуальне проектування та ER-діаграма", 2)
    add_text(doc, "На основі аналізу вимог було розроблено концептуальну модель бази даних.  Цей етап включає визначення основних сутностей, їх атрибутів та зв'язків без прив'язки до конкретної СУБД. Основні сутності та їх характеристики:")
    add_text(doc, "1.Users (Користувачі).  Це базова сутність, яка об'єднує всіх учасників системи. Вона містить атрибути автентифікації (email, password_hash), профілю (name, bio, avatar_url), ролі (role: client, performer, admin) та рейтингу (rating).  Також сюди входять поля для Telegram-інтеграції (telegram_id, telegram_username).")
    add_text(doc, "2.Services (Послуги). Сутність представляє одиницю товару на маркетплейсі. Вона пов'язана з користувачем-виконавцем (performer_id) та категорією (category_id).  Містить опис, базову ціну, терміни виконання, статус (active, paused) та медіа-матеріали (зображення, відео).")
    add_text(doc, "3.Orders (Замовлення). Центральна сутність, що фіксує факт економічної взаємодії.  Вона зв'язує клієнта (client_id), виконавця (performer_id) та послугу (service_id). Містить статус замовлення (pending, in_progress, completed, cancelled, disputed), фінальну ціну, дедлайн та JSON-поле для додаткових опцій (наприклад, «термінове виконання», «вихідні файли»).")
    add_text(doc, "4.  Messages (Повідомлення). Забезпечує комунікацію між учасниками.  Прив'язана до замовлення (order_id) та відправника (sender_id).  Має тип (текст/файл/система) та статус прочитання (is_read).")
    add_text(doc, "5.Reviews (Відгуки). Забезпечує систему репутації.  Зв'язує замовлення та користувачів. Важливо, що відгук можна залишити лише після успішного завершення замовлення.  Містить оцінку (1-5) та текстовий коментар.")
    add_text(doc, "6.Disputes (Диспути). Сутність для вирішення конфліктів. Зв'язана із замовленням та модератором.  Містить причину диспуту, опис проблеми та рішення адміністратора.")
    add_text(doc, "7.Policies (Політики).  Довідкова сутність для зберігання текстів юридичних документів (Terms of Service, Privacy Policy) різними мовами.")
    add_text(doc, "8.Categories (Категорії). Ієрархічний довідник послуг.  Містить назву, опис та посилання на батьківську категорію (parent_id), що дозволяє будувати деревоподібну структуру будь-якої глибини.")
    add_text(doc, "Зв'язки між сутностями реалізовані переважно як «один-до-багатьох» (1:M).  Наприклад, один користувач може мати багато замовлень, одне замовлення — багато повідомлень.  Зв'язок «багато-до-багатьох» (M:N) між послугами та тегами реалізовано через проміжну таблицю `service_tags`, що дозволяє ефективно фільтрувати послуги.")
    add_text(doc, "При проектуванні ER-діаграми використовувалася нотація Crow's Foot («вороняча лапка»), яка наочно відображає кардинальність зв'язків (обов'язковість та множинність).  Всі зв'язки є ідентифікуючими або неідентифікуючими залежно від бізнес-логіки.  Наприклад, повідомлення не може існувати без замовлення, тому це сильний (ідентифікуючий) зв'язок, і при видаленні замовлення повідомлення також повинні бути видалені.")

    add_heading(doc, "2.3. Логічне проектування та нормалізація", 2)
    add_text(doc, "Нормалізація — це процес організації даних у базі даних, який включає створення таблиць та встановлення зв'язків між ними відповідно до правил, спрямованих на захист даних та підвищення гнучкості бази даних шляхом усунення надлишковості та неузгоджених залежностей. Це дозволяє зменшити обсяг бази даних та спростити її підтримку.")
    add_text(doc, "У процесі проектування було послідовно досягнуто наступних нормальних форм:")
    add_text(doc, "Перша нормальна форма (1НФ) вимагає, щоб усі таблиці мали первинний ключ, а всі атрибути були атомарними (неподільними). Це означає, що в одній комірці не може зберігатися список значень (наприклад, список телефонів через кому) або група повторюваних полів (phone1, phone2).  У нашому проекті ми відмовилися від зберігання масивів ідентифікаторів у полях, замінивши їх на окремі таблиці зв'язків.  Використання JSONB є винятком, який допускається в сучасних РСУБД для напівструктурованих даних, де структура дійсно динамічна.")
    add_text(doc, "Друга нормальна форма (2НФ) вимагає, щоб таблиця знаходилася в 1НФ і всі неключові атрибути повністю залежали від повного первинного ключа. Це стосується таблиць зі складеними ключами.  У нашій схемі більшість таблиць мають простий сурогатний ключ (UUID), що автоматично спрощує дотримання 2НФ.  Наприклад, вартість замовлення фіксується в таблиці `orders` на момент створення і не залежить від поточної ціни послуги в таблиці `services`, яка може змінитися з часом.  Це забезпечує історичну достовірність даних.")
    add_text(doc, "Третя нормальна форма (3НФ) вимагає, щоб таблиця знаходилася в 2НФ і були відсутні транзитивні залежності. Атрибути, що не входять до первинного ключа, не повинні залежати від інших неключових атрибутів.  Наприклад, інформація про категорію послуги винесена в окрему таблицю `categories`, а в таблиці `services` зберігається лише посилання `category_id`.  Якби ми зберігали назву категорії прямо в таблиці послуг, це призвело б до дублювання даних і проблем при перейменуванні категорії (аномалія оновлення).")
    add_text(doc, "Нормальна форма Бойса-Кодда (BCNF) є посиленням 3НФ.  Вона вимагає, щоб кожна детермінанта була потенційним ключем.  У нашій схемі це дотримується.")
    add_text(doc, "Такий підхід дозволяє уникнути аномалій при вставці, оновленні та видаленні даних, забезпечуючи їх логічну цілісність.  Однак, в деяких випадках ми свідомо йдемо на контрольовану денормалізацію заради продуктивності. Наприклад, ми зберігаємо поточний рейтинг користувача в таблиці `users`, хоча його можна обчислити як середнє арифметичне всіх оцінок з таблиці `reviews`. Це дозволяє уникнути важких агрегатних запитів при кожному відображенні профілю користувача, що є критичним для високонавантаженої системи.  Оновлення цього поля відбувається через тригери.")

    # --- РОЗДІЛ 3 ---
    add_heading(doc, "РОЗДІЛ 3.  РЕАЛІЗАЦІЯ БАЗИ ДАНИХ СИСТЕМИ «HIWWER»", 1)
    
    add_heading(doc, "3.1.  Фізична реалізація структури даних", 2)
    add_text(doc, "Фізична модель бази даних реалізована у вигляді набору SQL-скриптів, які створюють структуру таблиць, індексів та обмежень. Для створення схеми використовується DDL (Data Definition Language).")
    add_text(doc, "Першим кроком є налаштування оточення та підключення необхідних розширень. Для генерації UUID використовується команда `CREATE EXTENSION IF NOT EXISTS \"uuid-ossp\";`.  Це дозволяє використовувати функцію `uuid_generate_v4()` як значення за замовчуванням для первинних ключів.UUID (Universally Unique Identifier) має ряд переваг над звичайними автоінкрементними ID (SERIAL): вони унікальні в межах усього кластеру, їх важче підібрати зловмисникам (security through obscurity), і вони дозволяють генерувати ID на стороні клієнта до вставки в БД, що спрощує роботу з транзакціями.")
    add_text(doc, "Створення таблиць відбувається у строго визначеному порядку для дотримання посилальної цілісності: спочатку створюються незалежні довідники (категорії, теги), потім основні сутності (користувачі, послуги), і в кінці — залежні сутності (замовлення, повідомлення, відгуки, диспути).")
    add_text(doc, "Розглянемо деталі реалізації ключових таблиць:")
    add_text(doc, "Таблиця `users`: встановлено унікальний індекс на поле `email` та `telegram_id`, що запобігає дублюванню облікових записів. Поле `role` обмежено списком допустимих значень через `CHECK constraint` (client, performer, admin).  Це гарантує, що в базу не потраплять некоректні дані навіть у випадку помилки в коді додатку.  Пароль зберігається у полі `password_hash` типу VARCHAR(255), що достатньо для зберігання хешів bcrypt.")
    add_text(doc, "Таблиця `orders`: використовує тип даних `TIMESTAMP WITH TIME ZONE` для полів `created_at` та `deadline`. Це критично важливо для міжнародного маркетплейсу, щоб коректно працювати з користувачами з різних часових поясів.  Поле `additional_options` має тип `JSONB` і значення за замовчуванням `'{}'::jsonb`, що гарантує наявність валідного JSON-об'єкта навіть якщо опції не обрані. Статус замовлення контролюється через ENUM або CHECK constraint.")
    add_text(doc, "Для оптимізації продуктивності створено індекси.  Окрім автоматичних індексів на Primary Key та Unique поля, створено індекси на поля зовнішніх ключів (`client_id`, `performer_id`, `service_id`) та поля, по яких часто відбувається фільтрація (`status`, `category_id`, `created_at`).  Це значно прискорює виконання запитів `JOIN` та `WHERE`.  PostgreSQL використовує B-Tree індекси за замовчуванням, які є ефективними для операцій рівності та діапазону.  Для полів JSONB використовуються GIN (Generalized Inverted Index) індекси, які дозволяють ефективно шукати ключі та значення всередині JSON-документів (наприклад, знайти всі замовлення з опцією 'urgent': true).")

    add_heading(doc, "3.2. Розробка серверної логіки (Тригери та Функції)", 2)
    add_text(doc, "Використання можливостей процедурної мови PL/pgSQL дозволило автоматизувати ряд рутинних операцій та забезпечити цілісність бізнес-даних безпосередньо на рівні СУБД. Це підвищує надійність системи, оскільки правила виконуються незалежно від того, який клієнт (веб-сайт, мобільний додаток, Telegram-бот) вносить зміни.")
    add_text(doc, "Автоматичний розрахунок рейтингу.  Рейтинг користувача є агрегованим показником, який впливає на довіру та ранжування у пошуку.  Замість того, щоб обчислювати його «на льоту» при кожному запиті профілю (що створює значне навантаження на БД при великій кількості відгуків), ми зберігаємо актуальне значення в полі `rating` таблиці `users`. Тригер `update_ratings` спрацьовує при події `INSERT`, `UPDATE` або `DELETE` в таблиці `reviews`.  Він обчислює нове середнє значення для відповідного користувача та оновлює запис у таблиці `users`. Це класичний приклад денормалізації заради продуктивності.")
    add_text(doc, "Аудит змін.  Для відстеження часу останньої активності та змін даних використовується універсальна функція та тригер `update_updated_at_column`. Він автоматично оновлює поле `updated_at` поточним часом (`NOW()`) при будь-якій зміні запису.  Це дозволяє легко реалізувати функціонал «нещодавно активні», «змінено нещодавно» та спрощує налагодження.")
    add_text(doc, "Валідація даних. Хоча основна валідація вхідних даних відбувається на рівні бекенду (API), додаткові перевірки на рівні БД (Defense in Depth) підвищують надійність системи. Наприклад, функція тригера перевіряє, чи `client_id` не дорівнює `performer_id` перед вставкою запису в таблицю замовлень або відгуків. Це унеможливлює ситуацію, коли користувач замовляє послугу сам у себе або накручує собі рейтинг, навіть якщо в коді API є помилка.")
    add_text(doc, "Також реалізовано функцію для автоматичного створення запису в таблиці `order_status_history` при зміні статусу замовлення. Це гарантує, що історія переходів станів буде повною і незмінною, що є критичним для вирішення диспутів.")

    add_heading(doc, "3.3. Реалізація специфічних модулів та інтеграцій", 2)
    
    add_subheading(doc, "3.3.1. Модуль історії статусів та чатів")
    add_text(doc, "Для забезпечення прозорості бізнес-процесів реалізовано повне логування історії замовлень.  Таблиця `order_status_history` діє як журнал подій (audit log).  Вона є незмінною (append-only) — записи в ній лише створюються, але ніколи не редагуються і не видаляються (окрім архівування).  Кожен запис містить посилання на замовлення, старий статус, новий статус, ID користувача, що ініціював зміну, та коментар. Це дозволяє відтворити хронологію подій у разі виникнення спорів.")
    add_text(doc, "Чат реалізовано через таблицю `messages`.  Для підтримки файлових вкладень використовується окрема таблиця `message_attachments`, що дозволяє прикріплювати до одного повідомлення декілька файлів.  Така структура є більш гнучкою, ніж зберігання масиву посилань у полі повідомлення.  Кожен файл має свій унікальний шлях зберігання (в хмарному сховищі S3 або локально), тип MIME та розмір, що дозволяє контролювати використання дискового простору та фільтрувати типи файлів.")

    add_subheading(doc, "3.3.2. Багатомовні політики")
    add_text(doc, "Система підтримує інтернаціоналізацію контенту (i18n).  Для юридичних документів (Terms of Service, Privacy Policy) створено таблицю `policies`.  Ключем у цій таблиці є складений унікальний індекс `(slug, language_code)`.  Це дозволяє отримувати текст документу потрібною мовою простим SQL-запитом, наприклад: `SELECT content FROM policies WHERE slug='terms' AND language_code='uk'`. Такий підхід дозволяє легко додавати нові мови без зміни структури БД та без необхідності перекомпіляції коду додатку.")

    add_subheading(doc, "3.3.3. Інтеграція з Telegram-ботом")
    add_text(doc, "Механізм прив'язки Telegram-акаунту реалізовано через тимчасові коди безпеки. Таблиця `telegram_linking_codes` зберігає згенерований випадковий код (наприклад, 6 цифр), ID користувача та час життя коду (expiration time, наприклад, 5 хвилин). Коли користувач вводить код у боті, відбувається пошук валідного запису.  Якщо код знайдено і він не прострочений, `telegram_id` записується в профіль користувача, а код видаляється. Це запобігає повторному використанню кодів та забезпечує безпеку процесу прив'язки.")
    add_text(doc, "Для реалізації сповіщень використовується механізм `LISTEN/NOTIFY` в PostgreSQL або періодичне опитування (polling) таблиці сповіщень. У нашому випадку ми використовуємо таблицю `notifications`, куди тригери або бекенд записують події (наприклад, «нове замовлення», «нове повідомлення»). Окремий сервіс бота періодично (наприклад, раз на секунду) вичитує нові непрочитані сповіщення (`is_read = false`) та відправляє їх користувачам у Telegram, після чого позначає як відправлені.  Це забезпечує асинхронність та надійність доставки.")

    add_subheading(doc, "3.3.4. Адміністративна панель та управління системою")
    add_text(doc, "Для ефективного управління платформою розроблено адміністративну панель з широким функціоналом. Основні компоненти адмін-панелі включають:")
    add_text(doc, "Dashboard (Панель керування).  Центральний елемент адмін-панелі, що надає агреговану інформацію про стан системи.  На дашборді відображаються ключові метрики: загальна кількість користувачів, виконавців, послуг та замовлень. Дані отримуються через оптимізовані SQL-запити з використанням функції `COUNT()`.  Для візуалізації використовуються графіки продажів по місяцях та розподіл замовлень по категоріях.  Таблиця `admin_settings` зберігає глобальні налаштування системи, такі як режим обслуговування (maintenance mode) та конфігурація дозволених типів файлів.")
    add_text(doc, "Управління користувачами та виконавцями.  Адміністратор має доступ до повного списку користувачів з можливістю фільтрації за роллю (client, performer, admin).  Реалізовано функціонал перегляду профілів, зміни статусу користувача та модерації контенту. Для виконавців додатково відображається статистика активності: кількість виконаних замовлень, середній рейтинг, кількість активних послуг.")
    add_text(doc, "Модерація замовлень та диспутів. Критично важливою функцією є система вирішення конфліктів.  Таблиця `disputes` зберігає інформацію про спори між клієнтами та виконавцями.  Адміністратор може бути призначений модератором диспуту (поле `moderator_id`). Для комунікації використовується окремий чат диспуту, реалізований через таблиці `dispute_messages` та `dispute_message_attachments`.  Адміністратор має доступ до повної історії замовлення, включаючи переписку, зміни статусів та вкладені файли, що дозволяє приймати обґрунтовані рішення.")
    add_text(doc, "Аналітика та звітність. Система генерує різноманітні звіти: продажі по місяцях (monthly sales), розподіл замовлень по категоріях, денна статистика продажів.  Використовуються агрегатні SQL-функції (`SUM`, `AVG`, `COUNT`) та групування даних (`GROUP BY`).  Функція `to_char` застосовується для форматування дат у зручний для читання формат.")
    add_text(doc, "Резервне копіювання та відновлення.  Таблиця `admin_settings` та окремі скрипти забезпечують автоматичне створення резервних копій бази даних.  Використовується утиліта `pg_dump` для створення логічних бекапів з можливістю відновлення на певний момент часу (Point-in-Time Recovery). Бекапи зберігаються у директорії `server/backups` з часовими мітками у форматі ISO 8601.")
    add_text(doc, "WebSocket моніторинг.  Таблиця `websocket_connections` дозволяє адміністратору відстежувати активні з'єднання користувачів в реальному часі. Це корисно для аналізу навантаження на систему та виявлення аномальної активності. Поля `connected_at`, `last_seen` та `is_active` забезпечують детальний контроль за сесіями.")

    add_subheading(doc, "3.3.5. Архітектура взаємодії Telegram-бота з базою даних")
    add_text(doc, "Telegram-бот є окремим сервісом, написаним на Python з використанням бібліотеки `python-telegram-bot`.  Взаємодія з базою даних відбувається через REST API бекенду, що забезпечує розділення відповідальності та можливість горизонтального масштабування.")
    add_text(doc, "Синхронізація користувачів.  При першому запуску бота користувач отримує пропозицію прив'язати свій Telegram-акаунт до профілю на платформі. Процес відбувається наступним чином: користувач отримує унікальний код на веб-сайті, який генерується бекендом та записується в таблицю `telegram_linking_codes` з обмеженим терміном дії (5 хвилин). Користувач вводить цей код у боті, бот відправляє запит до API для валідації.  При успішній валідації поля `telegram_id`, `telegram_username` та `telegram_chat_id` у таблиці `users` оновлюються, що встановлює зв'язок між обліковими записами.")
    add_text(doc, "Обробка сповіщень.  Сервіс сповіщень (NotificationService) працює в окремому асинхронному потоці. Він періодично опитує таблицю `notifications`, вибираючи записи з прапорцем `telegram_sent = FALSE`.  Для кожного сповіщення визначається тип події (new_order, message, status_change, deadline, dispute) та формується відповідне повідомлення на мові користувача (використовуючи поле `language_code` з таблиці `users`). Після успішної відправки в Telegram поля `telegram_sent` та `telegram_sent_at` оновлюються, що запобігає повторній відправці.  Такий підхід гарантує доставку навіть у випадку тимчасової недоступності Telegram API.")
    add_text(doc, "Обробка команд та callback-запитів.  Бот підтримує набір команд (/start, /help, /myorders, /language) та інтерактивні кнопки (inline keyboard). При отриманні команди бот звертається до API для отримання актуальних даних. Наприклад, команда /myorders викликає запит `GET /v1/orders? client_id={user_id}`, який повертає список замовлень користувача. Результат форматується та відображається у вигляді інтерактивних кнопок, кожна з яких містить callback_data з ідентифікатором замовлення.")
    add_text(doc, "Інтеграція з WebSocket для real-time оновлень. Хоча Telegram-бот не підтримує WebSocket напряму, він отримує миттєві оновлення через систему сповіщень.  Коли в чаті замовлення з'являється нове повідомлення, WebSocket-сервер бекенду створює запис у таблиці `notifications`, який автоматично підхоплюється сервісом сповіщень бота та доставляється користувачу.")

    add_subheading(doc, "3.3.6. Механізм резервного копіювання та міграції")
    add_text(doc, "Для забезпечення Disaster Recovery розроблено стратегію резервного копіювання.  Скрипт автоматичного бекапу використовує утиліту `pg_dump` для створення повного логічного знімку бази даних. Файли бекапів іменуються з використанням часової мітки (ISO 8601), стискаються (gzip) та відправляються на віддалений сервер або хмарне сховище. Скрипт запускається через системний планувальник завдань `cron` щоночі у період найменшого навантаження.")
    add_text(doc, "Керування змінами схеми БД здійснюється через міграції.  Кожна зміна (створення таблиці, додавання колонки, зміна типу даних) записується в окремий SQL-файл з порядковим номером та описом.  Це дозволяє автоматично застосовувати зміни при розгортанні нової версії додатку та гарантує, що структура БД на всіх серверах (development, staging, production) буде ідентичною.  Такий підхід (Database-as-Code) є стандартом у сучасній розробці та дозволяє уникнути помилок ручного адміністрування.")

    # --- ВИСНОВКИ ---
    add_heading(doc, "ВИСНОВКИ", 1)
    add_text(doc, "У ході виконання курсової роботи було проведено повний цикл проектування та реалізації реляційної бази даних для маркетплейсу цифрових послуг «Hiwwer».  Робота охопила всі етапи: від аналізу предметної області до фізичної реалізації та налаштування інфраструктури.")
    add_text(doc, "Основні результати роботи:")
    add_list_item(doc, "Проведено детальний аналіз предметної області електронної комерції послуг, виявлено специфічні вимоги до зберігання даних, такі як підтримка складних життєвих циклів замовлень, гнучкість атрибутів послуг та необхідність аудиту дій користувачів.")
    add_list_item(doc, "Обґрунтовано вибір СУБД PostgreSQL як оптимального рішення, що поєднує надійність реляційної моделі з гнучкістю роботи з JSON-даними, що є критичним для сучасних веб-додатків.")
    add_list_item(doc, "Спроектовано нормалізовану схему бази даних (3НФ), що складається з понад 15 взаємопов'язаних таблиць.  Схема забезпечує цілісність даних, відсутність надлишковості та підтримує всі необхідні бізнес-процеси (реєстрація, замовлення, комунікація, відгуки, диспути).")
    add_list_item(doc, "Реалізовано фізичну модель на базі СУБД PostgreSQL.  Використання сучасних можливостей цієї СУБД (JSONB для опцій, UUID для ключів, GIN індекси) дозволило створити гнучку, масштабовану та продуктивну систему.")
    add_list_item(doc, "Розроблено та впроваджено серверну логіку (тригери, функції PL/pgSQL), яка автоматизує критично важливі операції, такі як розрахунок рейтингів, валідація даних та аудит змін, знімаючи навантаження з прикладного рівня та гарантуючи узгодженість даних.")
    add_list_item(doc, "Успішно вирішено задачі інтеграції з зовнішніми системами (Telegram, AI) шляхом створення відповідних структур даних та механізмів синхронізації, що дозволяє створити омніканальний досвід для користувачів.")
    add_list_item(doc, "Забезпечено експлуатаційну надійність системи через впровадження процедур автоматичного резервного копіювання та міграції схеми, що мінімізує ризики втрати даних.")
    add_text(doc, "Практична цінність роботи полягає у створенні готового до використання компоненту даних, який може стати основою для повноцінного комерційного продукту. Отримана архітектура є модульною та масштабованою, що дозволяє подальший розвиток функціоналу без необхідності кардинальної перебудови структури даних.")
    add_text(doc, "У майбутньому планується розширити функціонал бази даних, додавши підтримку повнотекстового пошуку (Full Text Search) для покращення пошуку послуг, геопросторових даних (PostGIS) для пошуку виконавців поблизу, а також впровадити шардинг (секціонування) таблиць для горизонтального масштабування при значному зростанні навантаження.")

    # --- СПИСОК ДЖЕРЕЛ ---
    add_heading(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", 1)
    
    sources = [
        "PostgreSQL 16 Documentation [Електронний ресурс].  – Режим доступу: https://www.postgresql.org/docs/",
        "Date C.J.An Introduction to Database Systems / C.J.Date. – 8th Edition. – Addison-Wesley, 2003.  – 1328 p.",
        "Telegram Bot API Documentation [Електронний ресурс]. – Режим доступу: https://core.telegram.org/bots/api",
        "Конноллі Т.  Бази даних.  Проектування, реалізація та супровід.  Теорія та практика / Т. Конноллі, К. Бегг. – М.: Вільямс, 2003. – 1440 с.",
        "Google Cloud AI Documentation (Gemini API) [Електронний ресурс]. – Режим доступу: https://ai.google.dev/docs",
        "Fowler M.Patterns of Enterprise Application Architecture / M.Fowler. – Addison-Wesley Professional, 2002. – 560 p.",
        "Obe R.PostgreSQL: Up and Running: A Practical Guide to the Advanced Open Source Database / R.  Obe, L.Hsu. – O'Reilly Media, 2017. – 300 p.",
        "Richardson L.RESTful Web APIs / L.Richardson, M.Amundsen. – O'Reilly Media, 2013. – 406 p.",
        "PostgreSQL JSONB Guide [Електронний ресурс]. – Режим доступу: https://www.postgresqltutorial.com/postgresql-json/",
        "Hernandez M.J.Database Design for Mere Mortals: A Hands-On Guide to Relational Database Design / M.J.Hernandez.  – Addison-Wesley Professional, 2013. – 672 p.",
        "ISO/IEC 27001:2013 Information technology — Security techniques — Information security management systems — Requirements.",
        "GDPR - General Data Protection Regulation (EU) 2016/679 [Електронний ресурс]. – Режим доступу: https://gdpr-info.eu/"
    ]

    # Виводимо список джерел як нумерований список (без таблиці)
    for i, source in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(f"{i}. {source}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # --- ДОДАТКИ ---
    doc.add_page_break()
    
    # Додаток А
    p = doc.add_paragraph("ДОДАТОК А")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph("Фрагмент SQL-коду створення таблиць (schema.sql)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    code_schema = """-- Підключення розширень
CREATE EXTENSION IF NOT EXISTS "uuid-ossp";

-- Users Table (Таблиця користувачів)
CREATE TABLE users (
    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    name VARCHAR(255) NOT NULL,
    email VARCHAR(255) NOT NULL UNIQUE,
    password_hash VARCHAR(255) NOT NULL,
    role VARCHAR(50) NOT NULL CHECK (role IN ('client', 'performer', 'admin')),
    rating DECIMAL(3, 2) CHECK (rating >= 0 AND rating <= 5),
    bio TEXT,
    avatar_url VARCHAR(512),
    telegram_id VARCHAR(255) UNIQUE,
    telegram_username VARCHAR(255),
    telegram_chat_id VARCHAR(255),
    language_code VARCHAR(10) DEFAULT 'uk',
    is_performer BOOLEAN DEFAULT FALSE,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Service Categories (Категорії послуг)
CREATE TABLE service_categories (
    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    name VARCHAR(255) NOT NULL,
    slug VARCHAR(255) NOT NULL UNIQUE,
    description TEXT,
    icon VARCHAR(100),
    parent_id UUID REFERENCES service_categories(id),
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Services (Послуги)
CREATE TABLE services (
    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    performer_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    category_id UUID REFERENCES service_categories(id),
    title VARCHAR(255) NOT NULL,
    description TEXT,
    price DECIMAL(10, 2) NOT NULL CHECK (price >= 0),
    delivery_time INTEGER NOT NULL,
    rating DECIMAL(3, 2) DEFAULT 0,
    is_active BOOLEAN DEFAULT TRUE,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Orders (Замовлення)
CREATE TABLE orders (
    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    service_id UUID REFERENCES services(id),
    client_id UUID NOT NULL REFERENCES users(id),
    performer_id UUID REFERENCES users(id),
    category_id UUID REFERENCES service_categories(id),
    title VARCHAR(255) NOT NULL,
    requirements TEXT,
    status VARCHAR(50) NOT NULL CHECK (status IN (
        'pending', 'accepted', 'in_progress', 
        'delivered', 'completed', 'cancelled', 'disputed'
    )),
    price DECIMAL(10, 2) NOT NULL CHECK (price >= 0),
    deadline TIMESTAMP WITH TIME ZONE,
    additional_options JSONB NOT NULL DEFAULT '{}'::jsonb,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Messages (Повідомлення)
CREATE TABLE messages (
    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    order_id UUID NOT NULL REFERENCES orders(id) ON DELETE CASCADE,
    sender_id UUID REFERENCES users(id),
    receiver_id UUID REFERENCES users(id),
    content TEXT,
    message_type VARCHAR(50) DEFAULT 'text',
    is_read BOOLEAN DEFAULT FALSE,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Reviews (Відгуки)
CREATE TABLE reviews (
    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    order_id UUID NOT NULL REFERENCES orders(id),
    client_id UUID REFERENCES users(id),
    performer_id UUID REFERENCES users(id),
    reviewer_id UUID REFERENCES users(id),
    review_type VARCHAR(50) CHECK (review_type IN (
        'client_to_performer', 'performer_to_client'
    )),
    rating INTEGER NOT NULL CHECK (rating >= 1 AND rating <= 5),
    comment TEXT,
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

-- Індекси для оптимізації запитів
CREATE INDEX idx_orders_client_id ON orders(client_id);
CREATE INDEX idx_orders_performer_id ON orders(performer_id);
CREATE INDEX idx_orders_status ON orders(status);
CREATE INDEX idx_orders_created_at ON orders(created_at);
CREATE INDEX idx_messages_order_id ON messages(order_id);
CREATE INDEX idx_services_performer_id ON services(performer_id);
CREATE INDEX idx_services_category_id ON services(category_id);
CREATE INDEX idx_orders_additional_options ON orders USING GIN (additional_options);"""
    
    add_code_snippet(doc, code_schema)

    doc.add_page_break()

    # Додаток Б
    p = doc.add_paragraph("ДОДАТОК Б")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph("Фрагмент SQL-коду тригерів та функцій")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    code_triggers = """-- Функція автоматичного оновлення рейтингу виконавця
CREATE OR REPLACE FUNCTION update_performer_rating()
RETURNS TRIGGER AS $$
DECLARE
    avg_rating DECIMAL(3, 2);
BEGIN
    IF NEW.review_type = 'client_to_performer' THEN
        SELECT COALESCE(AVG(rating), 0) INTO avg_rating
        FROM reviews
        WHERE performer_id = NEW.performer_id
        AND review_type = 'client_to_performer';

        UPDATE users 
        SET rating = avg_rating,
            updated_at = NOW()
        WHERE id = NEW.performer_id;
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

-- Тригер для оновлення рейтингу
CREATE TRIGGER trigger_update_performer_rating
AFTER INSERT OR UPDATE OR DELETE ON reviews
FOR EACH ROW
EXECUTE FUNCTION update_performer_rating();

-- Функція автоматичного оновлення поля updated_at
CREATE OR REPLACE FUNCTION update_updated_at_column()
RETURNS TRIGGER AS $$
BEGIN
    NEW.updated_at = NOW();
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

-- Застосування тригера до всіх основних таблиць
CREATE TRIGGER trigger_users_updated_at
BEFORE UPDATE ON users
FOR EACH ROW EXECUTE FUNCTION update_updated_at_column();

CREATE TRIGGER trigger_orders_updated_at
BEFORE UPDATE ON orders
FOR EACH ROW EXECUTE FUNCTION update_updated_at_column();

CREATE TRIGGER trigger_services_updated_at
BEFORE UPDATE ON services
FOR EACH ROW EXECUTE FUNCTION update_updated_at_column();

-- Функція валідації: клієнт не може бути виконавцем
CREATE OR REPLACE FUNCTION validate_order_participants()
RETURNS TRIGGER AS $$
BEGIN
    IF NEW.client_id = NEW.performer_id THEN
        RAISE EXCEPTION 'Клієнт не може бути виконавцем власного замовлення';
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trigger_validate_order_participants
BEFORE INSERT OR UPDATE ON orders
FOR EACH ROW EXECUTE FUNCTION validate_order_participants();

-- Функція логування зміни статусу замовлення
CREATE OR REPLACE FUNCTION log_order_status_change()
RETURNS TRIGGER AS $$
BEGIN
    IF OLD.status IS DISTINCT FROM NEW.status THEN
        INSERT INTO order_status_history (
            order_id, old_status, new_status, changed_at
        ) VALUES (
            NEW.id, OLD.status, NEW.status, NOW()
        );
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trigger_log_order_status
AFTER UPDATE ON orders
FOR EACH ROW EXECUTE FUNCTION log_order_status_change();"""

    add_code_snippet(doc, code_triggers)

    doc.add_page_break()

    # Додаток В - ER-діаграма
    p = doc.add_paragraph("ДОДАТОК В")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph("Схема структури бази даних (ER-діаграма)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    ascii_schema = """
╔══════════════════════════════════════════════════════╗
║             СТРУКТУРА БАЗИ ДАНИХ HIWWER              ║
╚══════════════════════════════════════════════════════╝

┌──────────────────────┐    ┌──────────────────────┐
│  service_categories  │    │         tags         │
├──────────────────────┤    ├──────────────────────┤
│ PK id (UUID)         │    │ PK id (UUID)         │
│    name              │    │    name              │
│    slug (UNIQUE)     │    │    slug (UNIQUE)     │
│    description       │    └──────────┬───────────┘
│    icon              │               │ M:N
│ FK parent_id         │    ┌──────────┴───────────┐
└──────────┬───────────┘    │     service_tags     │
           │ 1:M            ├──────────────────────┤
           │                │ FK service_id        │
┌──────────┴───────────┐    │ FK tag_id            │
│       services       │◄───└──────────────────────┘
├──────────────────────┤
│ PK id (UUID)         │
│ FK performer_id      │──────────────────┐
│ FK category_id       │                  │
│    title             │                  │
│    description       │                  │
│    price             │                  │
│    delivery_time     │                  │
│    rating            │                  │
│    is_active         │                  │
└──────────┬───────────┘                  │
           │ 1:M                          │
┌──────────┴───────────┐                  │
│    service_images    │                  │
├──────────────────────┤                  │
│ PK id (UUID)         │                  │
│ FK service_id        │                  │
│    image_url         │                  │
│    position          │                  │
└──────────────────────┘                  │
                                          │
                         ┌────────────────┴────────────────┐
                         │              users              │
                         ├─────────────────────────────────┤
                         │ PK id (UUID)                    │
                         │    name                         │
                         │    email (UNIQUE)               │
                         │    password_hash                │
                         │    role (CHECK)                 │
                         │    rating                       │
                         │    telegram_id (UNIQUE)         │
                         │    telegram_username            │
                         │    language_code                │
                         │    is_performer                 │
                         └────────┬───────────┬────────────┘
                                  │           │
                     client_id    │           │  performer_id
                                  │           │
                         ┌────────┴───────────┴────────────┐
                         │             orders              │
                         ├─────────────────────────────────┤
                         │ PK id (UUID)                    │
                         │ FK service_id                   │
                         │ FK client_id                    │
                         │ FK performer_id                 │
                         │ FK category_id                  │
                         │    title                        │
                         │    requirements                 │
                         │    status (CHECK)               │
                         │    price                        │
                         │    deadline                     │
                         │    additional_options (JSONB)   │
                         └──┬────────┬────────┬────────────┘
                            │        │        │
          ┌─────────────────┘        │        └─────────────────┐
          │                          │                          │
          ▼ 1:M                     ▼ 1:M                     ▼ 1:M
┌────────────────────┐   ┌────────────────────┐   ┌────────────────────┐
│     messages       │   │order_status_history│   │      reviews       │
├────────────────────┤   ├────────────────────┤   ├────────────────────┤
│ PK id (UUID)       │   │ PK id (UUID)       │   │ PK id (UUID)       │
│ FK order_id        │   │ FK order_id        │   │ FK order_id        │
│ FK sender_id       │   │    old_status      │   │ FK client_id       │
│ FK receiver_id     │   │    new_status      │   │ FK performer_id    │
│    content         │   │    changed_at      │   │ FK reviewer_id     │
│    message_type    │   │    changed_by      │   │    review_type     │
│    is_read         │   └────────────────────┘   │    rating (1-5)    │
└─────────┬──────────┘                            │    comment         │
          │ 1:M                                   └────────────────────┘
          ▼
┌─────────────────────┐
│ message_attachments │
├─────────────────────┤
│ PK id (UUID)        │
│ FK message_id       │
│    file_url         │
│    file_name        │
│    file_size        │
│    mime_type        │
└─────────────────────┘

┌─────────────────────┐    ┌─────────────────────┐
│      disputes       │    │    notifications    │
├─────────────────────┤    ├─────────────────────┤
│ PK id (UUID)        │    │ PK id (UUID)        │
│ FK order_id         │    │ FK user_id          │
│ FK client_id        │    │    type             │
│ FK performer_id     │    │    content          │
│ FK moderator_id     │    │    is_read          │
│    reason           │    │    related_id       │
│    description      │    │    telegram_sent    │
│    status           │    │    telegram_sent_at │
│    resolution       │    └─────────────────────┘
└─────────┬───────────┘
          │ 1:M            ┌─────────────────────┐
          ▼                │telegram_linking_codes│
┌─────────────────────┐    ├─────────────────────┤
│   dispute_messages  │    │ PK id (UUID)        │
├─────────────────────┤    │ FK user_id          │
│ FK dispute_id       │    │    code (UNIQUE)    │
│ FK sender_id        │    │    expires_at       │
│    content          │    └─────────────────────┘
└─────────┬───────────┘    ┌─────────────────────┐
          │ 1:M            │   ai_chat_history   │
          ▼                ├─────────────────────┤
┌─────────────────────┐    │ PK id (UUID)        │
│dispute_msg_attachmts│    │    session_id       │
├─────────────────────┤    │    role             │
│ PK id (UUID)        │    │    message          │
│ FK message_id       │    │    tokens_count     │
│    file_name        │    │    created_at       │
│    file_url         │    └─────────────────────┘
└─────────────────────┘

┌─────────────────────┐    ┌─────────────────────┐
│websocket_connections│    │      policies       │
├─────────────────────┤    ├─────────────────────┤
│ PK id (UUID)        │    │ PK id (UUID)        │
│ FK user_id          │    │    slug             │
│    socket_id        │    │    language_code    │
│    connected_at     │    │    title            │
│    last_seen        │    │    content          │
│    is_active        │    │    content_markdown │
└─────────────────────┘    │ UNIQUE(slug, lang)  │
                           └─────────────────────┘
┌─────────────────────┐
│   admin_settings    │
├─────────────────────┤
│ PK id (SERIAL)      │
│    site_name        │
│    maintenance_mode │
│    allowed_file_types│
│    max_file_size    │
│    commission_rate  │
└─────────────────────┘
"""
    
    add_code_snippet(doc, ascii_schema)

    doc.add_page_break()

    # Додаток Г - Приклади SQL-запитів
    p = doc.add_paragraph("ДОДАТОК Г")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph("Приклади типових SQL-запитів системи")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    code_queries = """-- 1. Отримання списку замовлень користувача з деталями
SELECT 
    o.id,
    o.title,
    o.status,
    o.price,
    o.deadline,
    o.created_at,
    s.title AS service_title,
    p.name AS performer_name,
    p.rating AS performer_rating,
    c.name AS category_name
FROM orders o
LEFT JOIN services s ON o.service_id = s.id
LEFT JOIN users p ON o.performer_id = p.id
LEFT JOIN service_categories c ON o.category_id = c.id
WHERE o.client_id = $1
ORDER BY o.created_at DESC
LIMIT 20 OFFSET $2;

-- 2. Пошук послуг за категорією та ціновим діапазоном
SELECT 
    s.id,
    s.title,
    s.description,
    s.price,
    s.delivery_time,
    s.rating,
    u.name AS performer_name,
    u.avatar_url,
    c.name AS category_name,
    (SELECT COUNT(*) FROM reviews r WHERE r.performer_id = s.performer_id) AS reviews_count
FROM services s
JOIN users u ON s.performer_id = u.id
JOIN service_categories c ON s.category_id = c.id
WHERE s.category_id = $1
    AND s.price BETWEEN $2 AND $3
    AND s.is_active = TRUE
ORDER BY s.rating DESC, s.created_at DESC;

-- 3.  Статистика продажів по місяцях для адмін-панелі
SELECT 
    TO_CHAR(o.created_at, 'YYYY-MM') AS month,
    COUNT(*) AS total_orders,
    SUM(o.price) AS total_revenue,
    AVG(o.price) AS avg_order_value,
    COUNT(DISTINCT o.client_id) AS unique_clients,
    COUNT(DISTINCT o.performer_id) AS active_performers
FROM orders o
WHERE o.status = 'completed'
    AND o.created_at >= NOW() - INTERVAL '12 months'
GROUP BY TO_CHAR(o.created_at, 'YYYY-MM')
ORDER BY month DESC;

-- 4. Отримання непрочитаних сповіщень для Telegram-бота
SELECT 
    n.id,
    n.type,
    n.content,
    n.related_id,
    n.created_at,
    u.telegram_chat_id,
    u.language_code
FROM notifications n
JOIN users u ON n.user_id = u.id
WHERE n.telegram_sent = FALSE
    AND u.telegram_chat_id IS NOT NULL
    AND n.created_at > NOW() - INTERVAL '24 hours'
ORDER BY n.created_at ASC
LIMIT 100;

-- 5.  Розрахунок рейтингу виконавця з деталізацією
SELECT 
    u.id,
    u.name,
    u.rating,
    COUNT(r.id) AS total_reviews,
    COUNT(CASE WHEN r.rating = 5 THEN 1 END) AS five_star,
    COUNT(CASE WHEN r.rating = 4 THEN 1 END) AS four_star,
    COUNT(CASE WHEN r.rating = 3 THEN 1 END) AS three_star,
    COUNT(CASE WHEN r.rating <= 2 THEN 1 END) AS low_rating,
    ROUND(AVG(r.rating)::numeric, 2) AS calculated_rating
FROM users u
LEFT JOIN reviews r ON u.id = r.performer_id 
    AND r.review_type = 'client_to_performer'
WHERE u.id = $1
GROUP BY u.id, u.name, u.rating;

-- 6.  Пошук по JSONB полю additional_options
SELECT 
    o.id,
    o.title,
    o.price,
    o.additional_options->>'urgency' AS urgency_level,
    o.additional_options->>'source_files' AS include_source
FROM orders o
WHERE o.additional_options @> '{"urgent": true}'::jsonb
    AND o.status = 'pending';

-- 7.  Історія статусів замовлення для диспуту
SELECT 
    osh.id,
    osh.old_status,
    osh.new_status,
    osh.changed_at,
    u.name AS changed_by_name,
    u.role AS changed_by_role
FROM order_status_history osh
LEFT JOIN users u ON osh.changed_by = u.id
WHERE osh.order_id = $1
ORDER BY osh.changed_at ASC;"""

    add_code_snippet(doc, code_queries)

    # Зберігаємо документ
    doc.save('Coursework_Hiwwer_DB_v8.docx')
    print("=" * 60)
    print("Курсова робота (версія 8) успішно згенерована!")
    print("Файл: Coursework_Hiwwer_DB_v8.docx")
    print("=" * 60)
    print("\nПокращення у цій версії:")
    print("✓ Додано сторінку ЗМІСТ після титульної")
    print("✓ Покращено титульну сторінку (додано університет, факультет)")
    print("✓ Виправлено нумерацію сторінок (титульна без номера)")
    print("✓ Покращено форматування заголовків")
    print("✓ Додано Додаток Г з прикладами SQL-запитів")
    print("✓ Покращено ER-діаграму в Додатку В")
    print("✓ Всі шрифти примусово встановлено Times New Roman")

if __name__ == "__main__":
    create_coursework()